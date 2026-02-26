"""
Arabic language utilities for text normalization, encoding detection,
document loading, and RTL-aware chunking.

Features:
    - Auto-detect encoding for Arabic files (UTF-8, Windows-1256, ISO-8859-6).
    - Unicode normalization: remove tashkeel (diacritics), normalize alef variants,
      ta marbuta, and final yaa.
    - Detect whether a text is primarily Arabic.
    - Load Arabic .txt, .pdf, and .docx files with correct encoding.
    - RTL-aware text display helpers.

Requirements:
    pip install python-docx pypdf camel-tools  (camel-tools optional, for advanced NLP)
"""

import logging
import re
import unicodedata
from pathlib import Path
from typing import List, Optional

from langchain_core.documents import Document

from src.config import get_settings

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Arabic Unicode ranges & patterns
# ---------------------------------------------------------------------------

# Arabic script: U+0600–U+06FF, Arabic Presentation Forms: U+FE70–U+FEFF
_ARABIC_PATTERN = re.compile(r"[\u0600-\u06FF\uFE70-\uFEFF]")

# Tashkeel (diacritics / harakat)
_TASHKEEL = re.compile(
    r"[\u064B\u064C\u064D\u064E\u064F\u0650\u0651\u0652\u0653\u0654\u0655\u0656\u0657\u0658"
    r"\u0659\u065A\u065B\u065C\u065D\u065E\u065F\u0670]"
)

# Alef variants → plain alef
_ALEF_VARIANTS = re.compile(r"[إأآا]")
# Ta marbuta → ha
_TA_MARBUTA = re.compile(r"ة")
# Final yaa variants → yaa
_ALEF_MAQSURA = re.compile(r"ى")


# ---------------------------------------------------------------------------
# Encoding detection
# ---------------------------------------------------------------------------

def detect_encoding(file_path: str) -> str:
    """
    Detect the encoding of a file by trying a prioritized list of Arabic encodings.

    Tries UTF-8 first (most modern files), then common legacy Arabic encodings.
    Falls back to latin-1 as a last resort (never raises).

    Args:
        file_path: Path to the file.

    Returns:
        The name of the detected encoding (e.g. "utf-8", "windows-1256").
    """
    cfg = get_settings()
    encodings_to_try = cfg.arabic_encodings + ["latin-1"]

    for enc in encodings_to_try:
        try:
            with open(file_path, "r", encoding=enc, errors="strict") as f:
                f.read(4096)  # Read a sample to validate
            logger.debug("Detected encoding '%s' for file: %s", enc, file_path)
            return enc
        except (UnicodeDecodeError, LookupError):
            continue

    logger.warning("Could not reliably detect encoding for %s; defaulting to utf-8.", file_path)
    return "utf-8"


# ---------------------------------------------------------------------------
# Arabic text normalization
# ---------------------------------------------------------------------------

def remove_tashkeel(text: str) -> str:
    """
    Strip diacritical marks (tashkeel / harakat) from Arabic text.

    These marks affect pronunciation but are usually absent in typed text.
    Removing them improves matching between query and document tokens.

    Args:
        text: Raw Arabic text possibly containing diacritics.

    Returns:
        Text with all tashkeel characters removed.
    """
    return _TASHKEEL.sub("", text)


def normalize_alef(text: str) -> str:
    """
    Normalize all alef variants (أ إ آ ا) to plain alef (ا).

    Args:
        text: Input Arabic text.

    Returns:
        Text with unified alef characters.
    """
    return _ALEF_VARIANTS.sub("ا", text)


def normalize_ta_marbuta(text: str) -> str:
    """
    Normalize ta marbuta (ة) to ha (ه) for better stemming compatibility.

    Args:
        text: Input Arabic text.

    Returns:
        Text with ta marbuta replaced.
    """
    return _TA_MARBUTA.sub("ه", text)


def normalize_alef_maqsura(text: str) -> str:
    """
    Normalize alef maqsura (ى) to yaa (ي).

    Args:
        text: Input Arabic text.

    Returns:
        Text with alef maqsura replaced by yaa.
    """
    return _ALEF_MAQSURA.sub("ي", text)


def normalize_arabic_text(text: str, aggressive: bool = False) -> str:
    """
    Apply full Arabic normalization pipeline.

    Steps applied:
        1. Unicode NFC normalization.
        2. Remove tashkeel.
        3. Normalize alef variants.
        4. Normalize alef maqsura.
        5. (Aggressive) Normalize ta marbuta.
        6. Collapse multiple whitespace.

    Args:
        text:       Raw Arabic input string.
        aggressive: If True, also normalize ta marbuta (may alter word meaning).

    Returns:
        A cleaned, normalized Arabic string.
    """
    text = unicodedata.normalize("NFC", text)
    text = remove_tashkeel(text)
    text = normalize_alef(text)
    text = normalize_alef_maqsura(text)
    if aggressive:
        text = normalize_ta_marbuta(text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


# ---------------------------------------------------------------------------
# Language detection
# ---------------------------------------------------------------------------

def is_arabic_text(text: str, threshold: float = 0.3) -> bool:
    """
    Determine if a text is primarily written in Arabic.

    Counts the proportion of Arabic characters vs. total non-whitespace chars.

    Args:
        text:      The text to analyze.
        threshold: Minimum fraction of Arabic characters to classify as Arabic.

    Returns:
        True if the text is predominantly Arabic.
    """
    if not text:
        return False
    non_ws = [c for c in text if not c.isspace()]
    if not non_ws:
        return False
    arabic_chars = sum(1 for c in non_ws if _ARABIC_PATTERN.match(c))
    ratio = arabic_chars / len(non_ws)
    return ratio >= threshold


def detect_language(text: str) -> str:
    """
    Simple language detection returning "ar", "en", or "mixed".

    Args:
        text: Text to analyze.

    Returns:
        "ar"    — predominantly Arabic
        "en"    — predominantly Latin/English
        "mixed" — significant presence of both scripts
    """
    if not text:
        return "en"
    non_ws = [c for c in text if not c.isspace()]
    arabic_chars = sum(1 for c in non_ws if _ARABIC_PATTERN.match(c))
    ratio = arabic_chars / max(len(non_ws), 1)

    if ratio >= 0.6:
        return "ar"
    if ratio >= 0.2:
        return "mixed"
    return "en"


# ---------------------------------------------------------------------------
# Arabic-aware document loaders
# ---------------------------------------------------------------------------

def load_arabic_text_file(file_path: str) -> Document:
    """
    Load a plain-text file with automatic Arabic encoding detection and normalization.

    Args:
        file_path: Path to the .txt file.

    Returns:
        A LangChain Document with normalized page_content and metadata.

    Raises:
        FileNotFoundError: If the file does not exist.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    encoding = detect_encoding(file_path)
    with open(file_path, "r", encoding=encoding, errors="replace") as f:
        raw_text = f.read()

    cfg = get_settings()
    text = normalize_arabic_text(raw_text) if cfg.normalize_arabic else raw_text
    lang = detect_language(text)

    logger.info(
        "Loaded text file '%s' | encoding=%s | language=%s | chars=%d",
        path.name, encoding, lang, len(text),
    )
    return Document(
        page_content=text,
        metadata={"source": str(path), "encoding": encoding, "language": lang},
    )


def load_arabic_pdf(file_path: str) -> List[Document]:
    """
    Load a PDF file with Arabic content, normalizing text extracted from each page.

    Uses pypdf for extraction and applies Arabic normalization post-extraction.

    Args:
        file_path: Path to the .pdf file.

    Returns:
        A list of Documents, one per PDF page.
    """
    from pypdf import PdfReader

    path = Path(file_path)
    reader = PdfReader(str(path))
    cfg = get_settings()
    docs: List[Document] = []

    for page_num, page in enumerate(reader.pages, start=1):
        raw = page.extract_text() or ""
        lang = detect_language(raw)
        text = normalize_arabic_text(raw) if (cfg.normalize_arabic and lang in ("ar", "mixed")) else raw

        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "page": page_num,
                    "language": lang,
                    "total_pages": len(reader.pages),
                },
            ))

    logger.info("Loaded PDF '%s': %d pages extracted.", path.name, len(docs))
    return docs


def load_arabic_docx(file_path: str) -> Document:
    """
    Load a Word (.docx) document preserving Arabic text and paragraph order.

    Args:
        file_path: Path to the .docx file.

    Returns:
        A single Document with the full document text.
    """
    from docx import Document as DocxDocument  # python-docx

    path = Path(file_path)
    doc = DocxDocument(str(path))
    cfg = get_settings()

    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    raw_text = "\n".join(paragraphs)
    lang = detect_language(raw_text)
    text = normalize_arabic_text(raw_text) if (cfg.normalize_arabic and lang in ("ar", "mixed")) else raw_text

    logger.info("Loaded DOCX '%s': %d paragraphs | language=%s.", path.name, len(paragraphs), lang)
    return Document(
        page_content=text,
        metadata={"source": str(path), "language": lang},
    )


def load_arabic_document(file_path: str) -> List[Document]:
    """
    Dispatch loader based on file extension for Arabic-aware loading.

    Supports: .txt, .pdf, .docx

    Args:
        file_path: Path to the document.

    Returns:
        A list of Document objects.
    """
    ext = Path(file_path).suffix.lower()
    if ext == ".txt":
        return [load_arabic_text_file(file_path)]
    if ext == ".pdf":
        return load_arabic_pdf(file_path)
    if ext == ".docx":
        return [load_arabic_docx(file_path)]
    logger.warning("No Arabic-specific loader for extension '%s'; using default.", ext)
    from langchain_community.document_loaders import UnstructuredFileLoader
    return UnstructuredFileLoader(file_path).load()


# ---------------------------------------------------------------------------
# Directory-level loader
# ---------------------------------------------------------------------------

def load_arabic_directory(directory: str) -> List[Document]:
    """
    Recursively load all supported documents from a directory using Arabic-aware loaders.

    Args:
        directory: Path to a directory containing Arabic documents.

    Returns:
        A flat list of Documents from all files found.
    """
    supported = {".txt", ".pdf", ".docx"}
    docs: List[Document] = []

    for path in Path(directory).rglob("*"):
        if path.suffix.lower() in supported:
            try:
                loaded = load_arabic_document(str(path))
                docs.extend(loaded)
            except Exception as exc:
                logger.error("Failed to load '%s': %s", path, exc)

    logger.info("Arabic directory loader: %d documents from '%s'.", len(docs), directory)
    return docs
